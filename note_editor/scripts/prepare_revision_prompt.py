from __future__ import annotations

import argparse
from pathlib import Path
from docx import Document


def read_md(path: Path | None) -> str:
    if path is None:
        return ""
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")
    return path.read_text(encoding="utf-8")


def read_docx(path: Path) -> str:
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    doc = Document(path)
    paragraphs: list[str] = []

    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            paragraphs.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" / ".join(cells))

    return "\n\n".join(paragraphs)


OUTPUT_FORMAT = """\
# 出力形式

以下の形式で出力してください。

## 1. 全体レビュー
- 記事全体の良い点
- 改善すべき点
- 発信者らしさが保たれているか

## 2. 校正方針
- どの観点を優先して直したか
- どの表現をあえて残したか
- どの表現を削った、または弱めたか

## 3. 修正版全文
- noteへ貼り付けられるMarkdown形式
- 見出し、段落、余白を整える
- 元の思想・熱量を壊さない

## 4. 人間確認ポイント
- 投稿前に本人が確認すべき点
- 科学的事実確認が必要な箇所
- 表現の強さを最終判断すべき箇所
"""

TAG_GUIDE = """\
## 本文中タグの扱い

本文中には、執筆者が校正AIへ意図を伝えるためのタグが含まれる場合があります。
これらのタグは、校正時の判断材料として扱い、原則として修正版本文にはタグごと残さないでください。

### [EDITOR_NOTE]...[/EDITOR_NOTE]

執筆者から校正AIへの編集指示・意図説明です。
記事本文ではありません。
この内容を踏まえて、構成・見出し・段落接続・表現の強弱を調整してください。
ただし、修正版本文には `[EDITOR_NOTE]` タグおよびその中身をそのまま残さないでください。

### [FOOT_NOTE]...[/FOOT_NOTE]

脚注として扱う補足情報です。
本文の流れを妨げない補足説明として、必要に応じて脚注・注釈・補足段落の形に整えてください。
修正版本文では `[FOOT_NOTE]` タグは外し、noteに貼り付けやすい脚注形式または補足説明として整形してください。
"""


def _build_header(revision_round: int | None, mode: str) -> list[str]:
    """モードに応じた冒頭指示文を返す。"""
    parts: list[str] = []
    label = "GPT" if mode == "gpt" else "Claude"
    parts.append(f"# {label}校正依頼")
    parts.append("")
    parts.append("以下の校正指針に従って、note記事を校正してください。")
    parts.append("目的は、文章を一般的に綺麗にすることではなく、発信者の思想・温度感・科学的誠実さを保ちながら、読者に届く文章へ整えることです。")
    parts.append("")
    if revision_round is not None:
        parts.append("## 今回の校正段階")
        parts.append("")
        parts.append(f"これは{revision_round}次校正です。")
        parts.append("初稿からの全面的な書き換えではなく、前回までの修正意図を尊重しながら、残っている読みにくさ、論理の飛び、表現の強弱、科学的誠実さを確認してください。")
        parts.append("")
    return parts


def build_prompt(
    article_text: str,
    policy_text: str,
    author_text: str = "",
    reader_text: str = "",
    revision_round: int | None = None,
    mode: str = "gpt",
) -> str:
    """
    mode='gpt'  : 全文埋め込み型（従来どおり。GPT / Gemini向け）
    mode='claude': 記事本文のみのslim版（Claude Code + CLAUDE.md 運用向け）
    """
    parts: list[str] = _build_header(revision_round, mode)
    parts.append(TAG_GUIDE)

    if mode == "gpt":
        # GPTモード：コンテキストをすべて埋め込む
        if author_text:
            parts.append("---")
            parts.append("")
            parts.append("# 発信者プロフィール")
            parts.append("")
            parts.append(author_text)

        if reader_text:
            parts.append("---")
            parts.append("")
            parts.append("# 読者ペルソナ")
            parts.append("")
            parts.append(reader_text)

        parts.append("---")
        parts.append("")
        parts.append("# 校正指針")
        parts.append("")
        parts.append(policy_text)

    else:
        # Claudeモード：CLAUDE.mdが自動参照するためコンテキスト埋め込みは不要
        parts.append("---")
        parts.append("")
        parts.append("> **Note:** 発信者プロフィール・読者ペルソナ・校正指針は")
        parts.append("> `CLAUDE.md` の指示に従い、各参照ファイルから読み込んでください。")
        parts.append("")

    parts.append("---")
    parts.append("")
    parts.append("# 記事本文")
    parts.append("")
    parts.append(article_text)
    parts.append("")
    parts.append("---")
    parts.append("")
    parts.append(OUTPUT_FORMAT)

    return "\n".join(parts)


def main() -> None:
    parser = argparse.ArgumentParser(
        description="新規note記事の校正用プロンプトを生成します。"
    )

    parser.add_argument("--article", required=True, help="新規記事のdocxファイル")
    parser.add_argument("--policy", required=True, help="editorial_policy.md")
    parser.add_argument("--output", required=True, help="出力するprompt.md")

    parser.add_argument("--author", default=None, help="author_character.md 任意")
    parser.add_argument("--reader", default=None, help="persona_character.md 任意")
    parser.add_argument("--revision-round", type=int, default=None, help="何次校正か 任意")
    parser.add_argument(
        "--mode",
        choices=["gpt", "claude"],
        default="gpt",
        help=(
            "出力モード。"
            "'gpt': 全コンテキスト埋め込み型（GPT/Gemini向け、デフォルト）。"
            "'claude': 記事本文のみのslim版（Claude Code + CLAUDE.md 運用向け）。"
        ),
    )

    args = parser.parse_args()

    article_path = Path(args.article).expanduser().resolve()
    policy_path = Path(args.policy).expanduser().resolve()
    output_path = Path(args.output).expanduser().resolve()

    author_path = Path(args.author).expanduser().resolve() if args.author else None
    reader_path = Path(args.reader).expanduser().resolve() if args.reader else None

    article_text = read_docx(article_path)
    policy_text = read_md(policy_path)
    author_text = read_md(author_path)
    reader_text = read_md(reader_path)

    prompt = build_prompt(
        article_text=article_text,
        policy_text=policy_text,
        author_text=author_text,
        reader_text=reader_text,
        revision_round=args.revision_round,
        mode=args.mode,
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(prompt, encoding="utf-8")

    print(f"[OK] Generated ({args.mode} mode): {output_path}")


if __name__ == "__main__":
    main()