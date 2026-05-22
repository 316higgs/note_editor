from __future__ import annotations

import argparse
import difflib
import json
import re
from dataclasses import dataclass, asdict
from pathlib import Path

from docx import Document


@dataclass
class VersionFile:
    path: str
    filename: str
    order: int
    label: str
    paragraphs: list[str]


@dataclass
class ParagraphDiff:
    change_type: str
    before_index: int | None
    after_index: int | None
    before: str
    after: str


@dataclass
class RevisionDiff:
    from_file: str
    to_file: str
    from_label: str
    to_label: str
    changes: list[ParagraphDiff]
    stats: dict[str, int]


def normalize_text(text: str) -> str:
    text = text.replace("\u3000", " ")
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def extract_docx_text(path: Path) -> list[str]:
    doc = Document(path)
    paragraphs: list[str] = []

    for p in doc.paragraphs:
        text = normalize_text(p.text)
        if text:
            paragraphs.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [normalize_text(cell.text) for cell in row.cells]
            cells = [c for c in cells if c]
            if cells:
                paragraphs.append(" / ".join(cells))

    return paragraphs


def parse_order_and_label(path: Path) -> tuple[int, str]:
    stem = path.stem

    # 例:
    # 01_初版
    # 02-一次校正
    # 03 一次修正
    # 004_投稿前確認
    m = re.match(r"^\s*(\d+)[_\-\s　]*(.*)$", stem)

    if not m:
        # 番号がない場合は最後尾扱い
        return 999999, stem

    order = int(m.group(1))
    label = m.group(2).strip()

    if not label:
        label = f"version_{order}"

    return order, label


def load_versions(article_dir: Path) -> list[VersionFile]:
    docx_dir = article_dir
    files = sorted(docx_dir.glob("*.docx"))

    if not files:
        source_docx_dir = article_dir / "source_docx"
        if source_docx_dir.is_dir():
            docx_dir = source_docx_dir
            files = sorted(docx_dir.glob("*.docx"))

    versions: list[VersionFile] = []

    for path in files:
        if path.name.startswith("~$"):
            continue

        order, label = parse_order_and_label(path)
        paragraphs = extract_docx_text(path)

        versions.append(
            VersionFile(
                path=str(path),
                filename=path.name,
                order=order,
                label=label,
                paragraphs=paragraphs,
            )
        )

    versions.sort(key=lambda x: (x.order, x.filename))
    return versions


def diff_paragraphs(
    before: list[str],
    after: list[str],
) -> tuple[list[ParagraphDiff], dict[str, int]]:
    matcher = difflib.SequenceMatcher(None, before, after, autojunk=False)

    changes: list[ParagraphDiff] = []
    stats = {
        "added": 0,
        "deleted": 0,
        "modified": 0,
        "unchanged": 0,
    }

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            stats["unchanged"] += i2 - i1
            continue

        if tag == "delete":
            for i in range(i1, i2):
                changes.append(
                    ParagraphDiff(
                        change_type="deleted",
                        before_index=i + 1,
                        after_index=None,
                        before=before[i],
                        after="",
                    )
                )
                stats["deleted"] += 1

        elif tag == "insert":
            for j in range(j1, j2):
                changes.append(
                    ParagraphDiff(
                        change_type="added",
                        before_index=None,
                        after_index=j + 1,
                        before="",
                        after=after[j],
                    )
                )
                stats["added"] += 1

        elif tag == "replace":
            before_block = before[i1:i2]
            after_block = after[j1:j2]
            max_len = max(len(before_block), len(after_block))

            for k in range(max_len):
                b = before_block[k] if k < len(before_block) else ""
                a = after_block[k] if k < len(after_block) else ""

                if b and a:
                    change_type = "modified"
                    stats["modified"] += 1
                elif b:
                    change_type = "deleted"
                    stats["deleted"] += 1
                else:
                    change_type = "added"
                    stats["added"] += 1

                changes.append(
                    ParagraphDiff(
                        change_type=change_type,
                        before_index=i1 + k + 1 if b else None,
                        after_index=j1 + k + 1 if a else None,
                        before=b,
                        after=a,
                    )
                )

    return changes, stats


def build_revision_diffs(versions: list[VersionFile]) -> list[RevisionDiff]:
    revision_diffs: list[RevisionDiff] = []

    for prev, curr in zip(versions, versions[1:]):
        changes, stats = diff_paragraphs(prev.paragraphs, curr.paragraphs)

        revision_diffs.append(
            RevisionDiff(
                from_file=prev.filename,
                to_file=curr.filename,
                from_label=prev.label,
                to_label=curr.label,
                changes=changes,
                stats=stats,
            )
        )

    return revision_diffs


def write_json(
    article_dir: Path,
    versions: list[VersionFile],
    revision_diffs: list[RevisionDiff],
) -> Path:
    output = {
        "article_dir": str(article_dir),
        "versions": [asdict(v) for v in versions],
        "revision_diffs": [
            {
                **asdict(d),
                "changes": [asdict(c) for c in d.changes],
            }
            for d in revision_diffs
        ],
    }

    output_path = article_dir / f"{article_dir.name}_revision_diff.json"
    output_path.write_text(
        json.dumps(output, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return output_path


def write_markdown(
    article_dir: Path,
    versions: list[VersionFile],
    revision_diffs: list[RevisionDiff],
) -> Path:
    lines: list[str] = []

    lines.append("# note記事 校正差分レポート")
    lines.append("")
    lines.append(f"対象フォルダ: `{article_dir.name}`")
    lines.append("")

    lines.append("## 版一覧")
    lines.append("")

    for v in versions:
        lines.append(f"- `{v.filename}`")
        lines.append(f"  - order: `{v.order}`")
        lines.append(f"  - label: `{v.label}`")
        lines.append(f"  - paragraphs: {len(v.paragraphs)}")

    lines.append("")

    lines.append("## Geminiへの分析依頼")
    lines.append("")
    lines.append(
        "以下は、同一note記事の複数バージョン間の差分です。"
        "ファイル名とlabelは各編集段階の意味を表しています。"
        "各段階でどのような校正判断が行われたかを分析してください。"
    )
    lines.append("")
    lines.append("特に以下を分析してください。")
    lines.append("")
    lines.append("- 各段階で何が変更されたか")
    lines.append("- 文章の論理構造がどう変化したか")
    lines.append("- 読者への伝わりやすさを高める修正")
    lines.append("- 科学的厳密性に関する修正")
    lines.append("- 思想性・信仰的表現・美意識の調整")
    lines.append("- 削除された表現と、その理由の推定")
    lines.append("- 残された表現と、その意図の推定")
    lines.append("- 筆者固有の編集方針")
    lines.append("")

    if not revision_diffs:
        lines.append("## 差分")
        lines.append("")
        lines.append("比較可能なバージョンがありません。")
        lines.append("")

    for idx, diff in enumerate(revision_diffs, start=1):
        lines.append("---")
        lines.append("")
        lines.append(
            f"## 差分 {idx}: `{diff.from_label}` → `{diff.to_label}`"
        )
        lines.append("")
        lines.append(f"- from_file: `{diff.from_file}`")
        lines.append(f"- to_file: `{diff.to_file}`")
        lines.append(f"- added: {diff.stats['added']}")
        lines.append(f"- deleted: {diff.stats['deleted']}")
        lines.append(f"- modified: {diff.stats['modified']}")
        lines.append(f"- unchanged paragraphs: {diff.stats['unchanged']}")
        lines.append("")

        if not diff.changes:
            lines.append("変更は検出されませんでした。")
            lines.append("")
            continue

        for n, c in enumerate(diff.changes, start=1):
            lines.append(f"### Change {n}: {c.change_type}")
            lines.append("")

            if c.before:
                lines.append(f"**Before** paragraph `{c.before_index}`")
                lines.append("")
                lines.append("> " + c.before.replace("\n", "\n> "))
                lines.append("")

            if c.after:
                lines.append(f"**After** paragraph `{c.after_index}`")
                lines.append("")
                lines.append("> " + c.after.replace("\n", "\n> "))
                lines.append("")

    output_path = article_dir / f"{article_dir.name}_revision_diff_for_gemini.md"
    output_path.write_text("\n".join(lines), encoding="utf-8")
    return output_path


def process_article_dir(article_dir: Path) -> None:
    versions = load_versions(article_dir)

    if not versions:
        print(f"[SKIP] docxなし: {article_dir}")
        return

    if len(versions) == 1:
        print(f"[WARN] 比較対象が1ファイルのみ: {article_dir}")

    revision_diffs = build_revision_diffs(versions)

    json_path = write_json(article_dir, versions, revision_diffs)
    markdown_path = write_markdown(article_dir, versions, revision_diffs)

    print(f"[OK] {article_dir}")
    print(f"  - {json_path.name}")
    print(f"  - {markdown_path.name}")


def main() -> None:
    parser = argparse.ArgumentParser(
        description="docx版管理ファイルからGemini分析用の差分レポートを生成します。"
    )
    parser.add_argument(
        "target",
        type=str,
        help="記事フォルダ、または複数記事フォルダを含む親フォルダ",
    )
    parser.add_argument(
        "--recursive",
        action="store_true",
        help="親フォルダ配下の各記事フォルダをまとめて処理する",
    )

    args = parser.parse_args()
    target = Path(args.target).expanduser().resolve()

    if not target.exists():
        raise FileNotFoundError(f"対象が存在しません: {target}")

    if args.recursive:
        article_dirs = sorted(p for p in target.iterdir() if p.is_dir())
        for article_dir in article_dirs:
            process_article_dir(article_dir)
    else:
        process_article_dir(target)


if __name__ == "__main__":
    main()
